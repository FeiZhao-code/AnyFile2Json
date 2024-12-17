import json
import re

from docx import Document

def is_merged_cell(cell):
    """检查单元格是否被垂直合并"""
    try:
        # 检查 vMerge 是否存在且其值为 'continue'
        if hasattr(cell._tc, 'vMerge') and cell._tc.vMerge is not None:
            if isinstance(cell._tc.vMerge, str):
                return cell._tc.vMerge == 'continue'
            elif hasattr(cell._tc.vMerge, 'val'):
                return cell._tc.vMerge.val == 'continue'
    except AttributeError:
        pass
    return False

def is_empty_table_by_ratio(table, empty_threshold=0.8):
    """通过计算“|  |”模式占总单元格数的比例来判断表格是否为空"""
    total_cells = 0
    empty_like_cells = 0
    
    for row in table.rows:
        for cell in row.cells:
            # 跳过被垂直合并的单元格
            if is_merged_cell(cell):
                continue
            
            # 统计总单元格数
            total_cells += 1
            
            # 检查单元格文本是否匹配“|  |”模式
            cell_text = cell.text.strip()
            if cell_text == " | " or not cell_text:
                empty_like_cells += 1
    
    # 计算“|  |”模式占总单元格数的比例
    if total_cells == 0:
        return True  # 如果表格没有任何单元格，认为它是空的
    
    empty_ratio = empty_like_cells / total_cells
    
    # 如果“|  |”模式的比例超过阈值，则认为表格是空的
    return empty_ratio >= empty_threshold

def print_table_with_merged_cells(table):
    """打印表格内容，考虑合并单元格"""
    output = []
    for row_idx, row in enumerate(table.rows, start=1):
        row_text = []
        for col_idx, cell in enumerate(row.cells, start=1):
            if is_merged_cell(cell):
                continue  # 跳过被垂直合并的后续单元格
            # 添加单元格文本
            row_text.append(cell.text.strip())
        
        if row_text:  # 只打印非空行
            output.append(f"{' | '.join(row_text)}")
            # print(f"  Row {row_idx}: {' | '.join(row_text)}")
    return output
            
def is_heading(paragraph):
    # 修改后的正则表达式
    heading_pattern = re.compile(
        r'^(?P<section>\d+(?:\.\d+)*)(?:\.|、|\s)*(?P<title>.{15,}?)(?<![，。])(?=\s*\(|$)', 
        re.UNICODE
    )
    
    match = heading_pattern.match(paragraph.text.strip())
    if match:
        section = match.group('section')
        title = match.group('title').strip()
        if len(title) > 15:
            print(f"Matched with longer title, pass, Section: {section}, Title: '{title}'")
            return False
        print(f"Matched: {paragraph.text} -> Section: {section}, Title: '{title}'")
        return True
    else:
        if "Heading" in paragraph.style.name \
            or "标题" in paragraph.style.name \
            or "章" in paragraph.style.name \
            or ("节" in paragraph.text and len(paragraph.text) < 20):
            # or ("章" in paragraph.text and len(paragraph.text) < 20)\
            
            return True
        return False

def print_document_content(doc, print_paragraph_only=False):
    """遍历并打印文档中的所有段落和表格内容，可以选择只打印标题"""
    
    # 初始化一个列表来存储标题和正文
    content_list = []
    current_heading = ''
    content = []
    # 打印段落
    print("Printing paragraphs:")
    for para in doc.paragraphs:
        if para.text.strip() == '':
            continue
        if is_heading(para):
            if content != []:
                content_list.append({'title': current_heading, 'content': content})
                print(f"Added content: {content}")
                current_heading = f'[{para.style.name}:{para.text}]'
                content = []
            else:
                current_heading += f'-[{para.style.name}:{para.text}]'
        else:
            content.append(para.text)
                
    if content != []:
        content_list.append({'title': current_heading, 'content': content})

    print("\nPrinting tables:")
    # 打印表格
    for table_idx, table in enumerate(doc.tables, start=1):
        if is_empty_table_by_ratio(table):
            print(f"Table {table_idx} is empty (based on '|  |' ratio), skipping.")
            continue
        
        content_list.append({'title': f"Table {table_idx}:", 
                             'content': print_table_with_merged_cells(table)})
    
    # 将内容列表转换为JSON字符串
    json_data = json.dumps(content_list, indent=4, ensure_ascii=False)
    return json_data

def docx_to_json(docx_path, output_file):
    # 打开现有文档
    doc = Document(docx_path)

    json_data = print_document_content(doc, print_paragraph_only=True)  # 只打印标题和表格
    
    return json_data